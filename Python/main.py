import sqlite3
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
import matplotlib.pyplot as plt
import openpyxl
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_database():
    conn = sqlite3.connect("products.db")
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS products (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            date TEXT,
            name TEXT,
            manufacturer TEXT,
            sales INTEGER,
            price REAL
        )
    """)
    conn.commit()
    conn.close()


def load_data():
    conn = sqlite3.connect("products.db")
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM products")
    data = cursor.fetchall()
    conn.close()
    return data


def add_data(date, name, manufacturer, sales, price):
    conn = sqlite3.connect("products.db")
    cursor = conn.cursor()
    cursor.execute("INSERT INTO products (date, name, manufacturer, sales, price) VALUES (?, ?, ?, ?, ?)",
                   (date, name, manufacturer, sales, price))
    conn.commit()
    conn.close()


def delete_data(record_id):
    conn = sqlite3.connect("products.db")
    cursor = conn.cursor()
    cursor.execute("DELETE FROM products WHERE id=?", (record_id,))
    conn.commit()
    conn.close()


def update_data(record_id, date, name, manufacturer, sales, price):
    conn = sqlite3.connect("products.db")
    cursor = conn.cursor()
    cursor.execute("UPDATE products SET date=?, name=?, manufacturer=?, sales=?, price=? WHERE id=?",
                   (date, name, manufacturer, sales, price, record_id))
    conn.commit()
    conn.close()


def add_record():
    date = entry_date.get()
    name = entry_name.get()
    manufacturer = entry_manufacturer.get()
    sales = int(entry_sales.get())
    price = float(entry_price.get().replace(',', '.'))

    add_data(date, name, manufacturer, sales, price)
    update_table()


def delete_record():
    selected_item = table.selection()[0]
    record_id = table.item(selected_item)["values"][0]
    delete_data(record_id)
    update_table()


def edit_record():
    selected_item = table.selection()[0]
    record_id = table.item(selected_item)["values"][0]

    date = entry_date.get()
    name = entry_name.get()
    manufacturer = entry_manufacturer.get()
    sales = int(entry_sales.get())
    price = float(entry_price.get().replace(',', '.'))

    update_data(record_id, date, name, manufacturer, sales, price)
    update_table()


def on_table_click(event):
    selected_item = table.selection()[0]
    selected_record = table.item(selected_item)["values"]

    entry_date.delete(0, tk.END)
    entry_date.insert(tk.END, selected_record[1])

    entry_name.delete(0, tk.END)
    entry_name.insert(tk.END, selected_record[2])

    entry_manufacturer.delete(0, tk.END)
    entry_manufacturer.insert(tk.END, selected_record[3])

    entry_sales.delete(0, tk.END)
    entry_sales.insert(tk.END, selected_record[4])

    entry_price.delete(0, tk.END)
    entry_price.insert(tk.END, f"{selected_record[5]:,.2f}")


def create_word_report():
    data = load_data()

    doc = docx.Document()

    doc.add_heading("Отчет о продажах", 0)

    # Рассчитываем итоговую сумму продаж
    total_sales_amount = sum(float(row[4]) * float(row[5]) for row in data)

    # Добавляем запись об итоговой сумме продаж
    total_sales_paragraph = doc.add_paragraph()
    total_sales_paragraph.add_run(f"Итоговая сумма продаж: {total_sales_amount:,.2f} т.")
    total_sales_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Рассчитываем итоговое количество продаж
    total_sales_quantity = sum(row[4] for row in data)

    # Добавляем запись об итоговом количестве продаж
    total_quantity_paragraph = doc.add_paragraph()
    total_quantity_paragraph.add_run(f"Итоговое количество продаж: {total_sales_quantity}")
    total_quantity_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Заголовки таблицы
    headers = ["ID", "Дата", "Наименование", "Производитель", "Кол-во продаж", "Цена (за единицу)"]
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header

    # Заполнение таблицы данными
    for row_data in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row_data[0])
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
        row_cells[3].text = row_data[3]
        row_cells[4].text = str(row_data[4])
        row_cells[5].text = f"{row_data[5]:,.2f} т."

    # Сохранение документа
    doc.save("Отчет о продажах.docx")
    messagebox.showinfo("Отчет", "Отчет о продажах успешно сохранен в файле 'Отчет о продажах.docx'.")


def create_sales_chart():
    data = load_data()

    names = [row[2] for row in data]

    quantities = [row[4] for row in data]

    # Создание и настройка диаграммы
    plt.figure()
    plt.bar(names, quantities)
    plt.title("Количество продаж")
    plt.xlabel("Наименование")
    plt.ylabel("Кол-во продаж")
    plt.xticks(rotation=45, ha='right')

    # Отображение диаграммы
    plt.show()


def update_table():
    table.delete(*table.get_children())

    data = load_data()
    for row in data:
        table.insert('', tk.END, values=row)


def load_excel_data():
    file_path = "Book1.xlsx"
    workbook = openpyxl.load_workbook(file_path)
    worksheet = workbook.active
    data = []

    for row in worksheet.iter_rows(min_row=2, values_only=True):
        data.append(row)

    return data


def load_excel():
    data = load_excel_data()
    for row in data:
        add_data(row[0], row[1], row[2], row[3], row[4])
    update_table()


def create_pie_chart():
    data = load_data()

    names = [row[2] for row in data]
    quantities = [row[4] for row in data]

    # Получение топ 10 товаров
    top_names = []
    top_quantities = []
    other_quantity = 0
    for name, quantity in sorted(zip(names, quantities), key=lambda x: x[1], reverse=True):
        if len(top_names) < 10:
            top_names.append(name)
            top_quantities.append(quantity)
        else:
            other_quantity += quantity

    # Добавление остальных товаров в список топовых
    top_names.append("Другие товары")
    top_quantities.append(other_quantity)

    # Создание диаграммы
    fig, ax = plt.subplots()
    ax.pie(top_quantities, labels=top_names, autopct='%1.1f%%', startangle=90)
    ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
    plt.title("Топ 10 товаров по количеству продаж")

    # Отображение диаграммы
    plt.show()


# Создание основного окна
main_window = tk.Tk()
main_window.title("Продажи товаров")

# Верхняя панель
top_frame = ttk.Frame(main_window)
top_frame.pack(side=tk.TOP, padx=10, pady=10)

ttk.Label(top_frame, text="Дата:").grid(row=0, column=0, padx=5, pady=5)
entry_date = ttk.Entry(top_frame)
entry_date.grid(row=0, column=1, padx=5, pady=5)

ttk.Label(top_frame, text="Наименование:").grid(row=0, column=2, padx=5, pady=5)
entry_name = ttk.Entry(top_frame)
entry_name.grid(row=0, column=3, padx=5, pady=5)

ttk.Label(top_frame, text="Производитель:").grid(row=0, column=4, padx=5, pady=5)
entry_manufacturer = ttk.Entry(top_frame)
entry_manufacturer.grid(row=0, column=5, padx=5, pady=5)

ttk.Label(top_frame, text="Кол-во продаж:").grid(row=0, column=6, padx=5, pady=5)
entry_sales = ttk.Entry(top_frame)
entry_sales.grid(row=0, column=7, padx=5, pady=5)

ttk.Label(top_frame, text="Цена (за единицу):").grid(row=0, column=8, padx=5, pady=5)
entry_price = ttk.Entry(top_frame)
entry_price.grid(row=0, column=9, padx=5, pady=5)

# Кнопки
buttons_frame = ttk.Frame(main_window)
buttons_frame.pack(side=tk.TOP, padx=10, pady=10)

ttk.Button(buttons_frame, text="Добавить", command=add_record).grid(row=0, column=0, padx=5, pady=5)
ttk.Button(buttons_frame, text="Удалить", command=delete_record).grid(row=0, column=1, padx=5, pady=5)
ttk.Button(buttons_frame, text="Изменить", command=edit_record).grid(row=0, column=2, padx=5, pady=5)
ttk.Button(buttons_frame, text="Сохранить отчет в Word", command=create_word_report).grid(row=0, column=3, padx=5,
                                                                                          pady=5)
ttk.Button(buttons_frame, text="Создать диаграмму", command=create_sales_chart).grid(row=0, column=4, padx=5, pady=5)
ttk.Button(buttons_frame, text="Загрузить из Excel", command=load_excel).grid(row=0, column=5, padx=5, pady=5)
ttk.Button(buttons_frame, text="Создать круговую диаграмму", command=create_pie_chart).grid(row=0, column=6, padx=5,
                                                                                            pady=5)

# Таблица с данными
table_frame = ttk.Frame(main_window)
table_frame.pack(side=tk.BOTTOM, padx=10, pady=10)

table = ttk.Treeview(table_frame, columns=("1", "2", "3", "4", "5", "6"), show="headings", height=10)

table.pack(side=tk.LEFT, padx=5, pady=5)

table.heading(1, text="ID")
table.heading(2, text="Дата")
table.heading(3, text="Наименование")
table.heading(4, text="Производитель")
table.heading(5, text="Кол-во продаж")
table.heading(6, text="Цена (за единицу)")

table.column(1, width=50, anchor='center')
table.column(2, width=100, anchor='center')
table.column(3, width=200, anchor='center')
table.column(4, width=150, anchor='center')
table.column(5, width=100, anchor='center')
table.column(6, width=150, anchor='center')

# Прокрутка для таблицы
scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=table.yview)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
table.configure(yscrollcommand=scrollbar.set)

# Событие клика по таблице
table.bind("<ButtonRelease-1>", on_table_click)

# Создание базы данных и загрузка данных
create_database()
update_table()

main_window.resizable(width=False, height=False)
main_window.iconbitmap("logo.ico")

main_window.mainloop()
